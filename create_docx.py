#!/usr/bin/env python3
"""Generate system documentation as .docx file - natural paragraph format"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title_run = title.add_run("System Documentation: Refine Data Studio")
title_run.font.size = Pt(24)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Data Cleaning and Analytics System (FEDM)")
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.bold = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# 1. Introduction
doc.add_heading('1. Introduction', level=1)

doc.add_heading('Brief Description', level=2)
doc.add_paragraph('Refine Data Studio is a web-based platform designed to help analysts and engineers work with messy datasets. Built with Python and Streamlit, the system provides a straightforward interface where users can upload raw data, understand what problems exist in it, apply cleaning transformations, and then explore the results through statistical analysis and interactive charts. The platform handles the repetitive work of data preparation, making it possible to move from raw CSV files to analysis-ready datasets in a matter of minutes rather than hours.')

doc.add_heading('Purpose of the Project', level=2)
doc.add_paragraph('The real problem this system solves is the gap between data collection and data analysis. When data arrives from production systems, APIs, or manual entry, it almost never comes in perfect condition. There are missing values, formatting inconsistencies, duplicate records, type mismatches, and invalid entries. In a typical workflow, analysts spend 30 to 40 percent of their time just preparing data for analysis instead of actually analyzing it. This system automates the detection and fixing of common data quality issues, which means analysts can focus on the insights rather than the mechanics of data preparation.')

doc.add_heading('Target Users', level=2)
doc.add_paragraph('The system is designed for several different types of users. Data analysts who regularly work with datasets and need to clean them before building reports or models would find this helpful. Business analysts who want to understand the quality of data and spot trends would benefit from the automatic insights. Database administrators need to validate data when importing from external sources or cleaning up legacy systems. And for students or researchers, the system provides a practical way to learn data cleaning techniques without having to write code from scratch.')

# 2. Problem Statement
doc.add_page_break()
doc.add_heading('2. Problem Statement', level=1)

doc.add_heading('What Data Problem Are You Solving?', level=2)
doc.add_paragraph('Real data is messy. When you download a CSV file from a database or receive an Excel spreadsheet from a colleague, you inevitably encounter a mix of problems. Some cells are empty or contain NULL values where there should be data. Other fields have been filled in inconsistently—maybe dates are stored in four different formats, or department names appear as "SALES", "Sales", and "sales" in different rows. Duplicate records pop up when data has been merged from multiple sources. Numbers might be stored as text strings, making it impossible to calculate averages or sums. Values exist that don\'t make sense in context, like negative salaries or ages over 150. Text fields have extra spaces at the beginning or end. And sometimes you\'ll find an obvious outlier like a salary of nine million dollars in a dataset where most salaries are between thirty and two hundred thousand.')

doc.add_heading('Why Is It Important?', level=2)
doc.add_paragraph('Bad data leads to bad decisions. If you build a report or statistical model on data that contains errors and inconsistencies, your results will be wrong. The problems compound when multiple people are working with the same dataset and each one has to spend time figuring out and fixing the same issues. Beyond just wasting time, there\'s a compliance aspect. In regulated industries, you need to be able to demonstrate that your data has been validated and cleaned according to documented procedures. Having a systematic approach to data cleaning means you can reproduce your work, audit what was changed and why, and trust that your analysis is built on a solid foundation.')

# 3. System Overview
doc.add_page_break()
doc.add_heading('3. System Overview', level=1)

doc.add_heading('General Description', level=2)
doc.add_paragraph('Refine is built around a simple idea: take data in one format, let users see what\'s wrong with it, give them tools to fix common problems, and then let them explore the cleaned data. The system combines several different layers working together. On the front end, there\'s a web interface built with Streamlit that runs right in a browser. The data processing happens with Pandas, a Python library that makes it easy to manipulate rows and columns. The statistical analysis uses standard methods to find patterns and outliers. Charts are built with Plotly, which gives you interactive visualizations you can zoom into or export as images. And everything is backed by a MySQL database so that when you clean a dataset, the results are saved and you can come back to them later.')

doc.add_heading('Key Features', level=2)
doc.add_paragraph('The system has six main sections that users move through in order. Upload is where you bring in your CSV or Excel file and the system reads it into memory. Profile gives you a statistical overview of your dataset so you understand how many rows and columns you have, what data types exist, where values are missing, and what the range and distribution look like for numeric columns. Clean is where you apply transformations. You can fill in missing values using the median of a column, remove duplicate rows, convert text columns to numbers, standardize date formats or text capitalization, and remove values that fall outside expected ranges. Compare shows you the before and after side by side so you can verify that the cleaning did what you expected. Insights automatically generates statistical findings about your data including correlations between columns. And Dashboard lets you build interactive charts from your cleaned data so you can explore patterns and share results.')

doc.add_heading('Technology Stack', level=2)
doc.add_paragraph('The system is written in Python 3.11. The web interface uses Streamlit, which is designed specifically for building data applications quickly without needing to write HTML or JavaScript. Data manipulation happens with Pandas and NumPy. Statistical analysis uses SciPy and scikit-learn. Charts are built with Plotly, which is one of the few charting libraries that makes it easy to create truly interactive visualizations. Data is persisted in MySQL 5.7 or later using SQLAlchemy as an object-relational mapper, which means the Python code doesn\'t have to worry about writing raw SQL.')

# 4. System Workflow
doc.add_page_break()
doc.add_heading('4. System Workflow', level=1)

doc.add_paragraph('The data moves through six stages: Upload, Profile, Clean, Compare, Insights, and Dashboard. At each stage, the user sees results and decides what to do next.')

doc.add_heading('Step 1: Upload', level=2)
doc.add_paragraph('A user opens the application in their browser and the first thing they see is the Upload section. They can drag a CSV or Excel file into the drop zone or click to browse their file system. The system detects the file format automatically and reads it into memory as a DataFrame, which is Pandas\' way of representing a table of data. Once the file is loaded, some basic information is stored in the database: the name of the file, when it was uploaded, and how many rows and columns it has.')

doc.add_heading('Step 2: Profile', level=2)
doc.add_paragraph('When you click Profile, the system analyzes your data column by column. For each column, it figures out what type of data it contains. Are these numbers, text, dates, or something else? It counts how many cells are empty or NULL. It figures out how many unique values exist in the column. For numeric columns, it calculates the mean, median, minimum, maximum, and standard deviation so you understand the shape of the distribution. For text columns, it shows you the five most common values and how often each appears. This profile gives you a complete picture of what you\'re working with.')

doc.add_heading('Step 3: Clean', level=2)
doc.add_paragraph('The Clean section is where the actual transformation happens. The system provides six different operations you can apply. You can fill missing values using various strategies like the median of a column or the most common value. You can remove rows that are exact duplicates of other rows. You can convert a column from one data type to another, like turning text that represents numbers into actual numeric types. You can standardize text by removing extra spaces, converting to uppercase or lowercase, or cleaning up special characters. You can reformat dates so they all follow the same pattern. And you can filter out invalid values by removing anything that falls outside a specified range or doesn\'t match a pattern. Every time you apply an operation, it\'s logged to the database with the parameters you used, so you have a record of exactly what was changed.')

doc.add_heading('Step 4: Compare', level=2)
doc.add_paragraph('After you\'ve applied some cleaning operations, the Compare section shows you the original data and the cleaned data side by side. Cells that have been changed are highlighted in yellow. Cells that were filled in are highlighted in green. Rows that were removed are shown in red. This visual comparison makes it easy to verify that your cleaning operations did what you intended and didn\'t accidentally break something.')

doc.add_heading('Step 5: Insights', level=2)
doc.add_paragraph('Once your data is clean, the Insights section automatically generates findings. For numeric columns it shows you the mean and median values, the range of data, and identifies any outliers that are extremely far from the typical values. For categorical columns it shows you what values exist, how common each one is, and how much diversity exists in the column. The system also calculates correlations between pairs of numeric columns, which tells you whether two variables move together or independently. All of this is displayed both as text that you can read and as visualizations like a heatmap showing correlation strength.')

doc.add_heading('Step 6: Dashboard', level=2)
doc.add_paragraph('The final section is Dashboard. Here you can create custom visualizations from your cleaned data. You pick what type of chart you want, which columns to use for the x-axis and y-axis, and whether you want to aggregate the data using sum, mean, or count. Plotly renders the chart interactively so you can hover to see exact values, zoom in, pan around, and toggle series on and off using the legend. You can download any chart as a PNG image. The charts are saved to the database so you can close the application and come back later to see the same visualizations.')

# 5. Data Profiling
doc.add_page_break()
doc.add_heading('5. Data Profiling', level=1)

doc.add_heading('How the System Analyzes Your Data', level=2)
doc.add_paragraph('When you upload a file and click Profile, the system runs a statistical analysis on every column. This happens automatically and takes just a few seconds. The goal is to understand what you have so that you know what needs to be fixed.')

doc.add_heading('What Gets Detected', level=2)
doc.add_paragraph('The system figures out the data type of each column by looking at the actual values. If a column contains numbers like 50000 or 65.5, it classifies that as numeric. If it contains text, it marks it as categorical. If it recognizes date patterns like "2021-03-15" or "March 15 2021", it identifies it as datetime. For each column, it counts how many cells are empty or contain NULL values and calculates what percentage that represents. It counts the number of unique distinct values in the column, which gives you a sense of the diversity of the data.')

doc.add_paragraph('For numeric columns, the profiler calculates standard statistical measures. The mean is the average of all values. The median is the middle value when you sort everything from smallest to largest. The standard deviation tells you how spread out the values are, with higher numbers meaning the values are more scattered. Minimum and maximum show you the range of data. All of these statistics together give you a sense of what typical values look like and where outliers might exist.')

doc.add_paragraph('For categorical or text columns, the system takes a different approach. It shows you the five most common values and how many times each appears. This helps you understand what the column is actually used for and whether there are unexpected variations. For example, if you have a status column that should only contain "Active" and "Inactive", but the profile shows you values like "Active", "ACTIVE", "active", and "Inactive", you\'ve immediately spotted a formatting problem that needs fixing.')

# 6. Data Cleaning Methods
doc.add_page_break()
doc.add_heading('6. Data Cleaning Methods', level=1)

doc.add_paragraph('The system provides six different cleaning operations. Each one addresses a specific type of data quality problem. You can apply them in any order and as many times as you want.')

doc.add_heading('Fill Missing Values', level=2)
doc.add_paragraph('Missing values are NULL or empty cells in your dataset. The system gives you several options for how to handle them. You can drop the entire row, which makes sense if missing data is rare and the rest of the row is not important. You can drop the entire column if it has so many missing values that it\'s not useful. You can fill numeric columns with the mean or median, which preserves the overall distribution of values. You can fill categorical columns with the mode, which is the most common value. You can provide a custom value like "Unknown" that you want used everywhere. Or you can use forward fill or backward fill, which are useful for time-series data where you want to propagate a value forward or backward in sequence.')

doc.add_heading('Drop Duplicates', level=2)
doc.add_paragraph('Sometimes your dataset contains rows that are exact duplicates of other rows. This happens when data has been merged from multiple sources or when records were accidentally imported twice. The system can remove these duplicates automatically. You can choose to remove rows that are identical across all columns, or you can specify that you only care about duplicates in certain columns. For instance, if your employee table somehow has two records with the same ID number, you might want to remove one of them.')

doc.add_heading('Convert Type', level=2)
doc.add_paragraph('Often data arrives in the wrong format. A column of salaries might be stored as text instead of numbers, which prevents you from calculating totals or averages. Dates might be stored as text strings. IDs might be floats when they should be integers. The conversion operation lets you change a column from one type to another. The system is smart about it too, removing currency symbols, stripping whitespace, and handling common formatting variations. If a value can\'t be converted, it\'s flagged so you can decide what to do with it.')

doc.add_heading('Standardize Text', level=2)
doc.add_paragraph('Text data often comes in inconsistent formats. Department names might appear as "SALES", "Sales", and "sales". Status fields might have "Active" and "ACTIVE". Names might have leading or trailing spaces. The standardize operation lets you apply transformations like converting everything to uppercase or lowercase, removing extra whitespace, converting to title case, or removing special characters. This ensures that identical values are truly identical, which makes analysis and grouping much cleaner.')

doc.add_heading('Standardize Dates', level=2)
doc.add_paragraph('Dates are famously messy because different countries and systems use different formats. You might have some dates formatted as 2021-03-15, others as 03/15/2021, and others as "March 15 2021". The system can parse all of these variations and convert them to a single standard format: YYYY-MM-DD. It also detects when dates are invalid, like "2021-13-45" or "not-a-date", and can flag or remove those rows.')

doc.add_heading('Filter Invalid', level=2)
doc.add_paragraph('This operation removes rows where a column contains values that don\'t make sense. You specify a range, for example salaries must be between 25,000 and 500,000, and any row with a value outside that range gets removed. You can also provide a list of allowed values, like a status field that should only ever contain "Active", "Inactive", or "Pending". Or you can provide a regular expression pattern for more complex validation, like an email column that must follow email format rules.')

# 7. Insights Generation
doc.add_page_break()
doc.add_heading('7. Insights Generation', level=1)

doc.add_paragraph('Once your data is clean, the system automatically generates insights about it. These are findings that help you understand what the data is telling you.')

doc.add_heading('Numeric Analysis', level=2)
doc.add_paragraph('For columns with numbers, the system calculates basic statistics. It shows the mean and median so you understand central tendency. It calculates the range from minimum to maximum to show how spread out the data is. It identifies outliers, which are values that are extremely far from the typical range, usually defined as more than three standard deviations away from the mean. These might be data errors or they might be real but unusual cases that need investigation.')

doc.add_heading('Categorical Analysis', level=2)
doc.add_paragraph('For text columns, the system shows you the distribution of values. It displays the most common values and what percentage of rows each represents. It shows how much diversity exists in the column. And it reports what percentage of values are missing. This gives you a sense of whether the column is useful for analysis and whether the data is dominated by a few common values or more evenly distributed.')

doc.add_heading('Correlation Analysis', level=2)
doc.add_paragraph('The system calculates correlations between pairs of numeric columns. Correlation is a number between minus one and plus one that tells you whether two variables move together. A correlation near one means they move in the same direction: when one goes up, the other tends to go up. A correlation near minus one means they move in opposite directions. A correlation near zero means they\'re independent. This is useful for spotting relationships in your data. For instance, in an employee dataset, you\'d expect a strong correlation between years of experience and salary. If that correlation is missing, it might indicate a problem with how the data was entered.')

# 8. Dashboard & Visualization
doc.add_page_break()
doc.add_heading('8. Dashboard & Visualization', level=1)

doc.add_paragraph('The Dashboard section lets you build interactive charts from your cleaned data. These aren\'t static images but real interactive visualizations that you can explore.')

doc.add_heading('Available Chart Types', level=2)
doc.add_paragraph('Bar charts are useful when you want to compare values across categories. For instance, you might want to see average salary by department. Line charts are good for showing trends over time. Pie charts show composition, displaying what percentage of the total each category represents. Scatter plots reveal relationships between two numeric variables, making it easy to spot correlations visually. Histograms show the distribution of a single numeric column, displaying how many values fall into each range.')

doc.add_heading('How to Build a Chart', level=2)
doc.add_paragraph('To create a chart, you select the chart type you want, choose which column goes on the x-axis, optionally choose a column for the y-axis, and specify how to aggregate the data. If you\'re creating a bar chart showing average salary by department, the x-axis would be department, the y-axis would be salary, and the aggregation would be mean. The system automatically groups the data and calculates the values needed for the chart.')

doc.add_heading('Interactive Features', level=2)
doc.add_paragraph('All charts are built with Plotly, a library that makes visualizations interactive. When you hover your mouse over a bar, line point, or pie slice, you see the exact value. You can click and drag to zoom into a region you\'re interested in, and the axes rescale automatically. You can pan around when you\'re zoomed in. The legend on the side lets you click to hide or show specific series. And you can download the chart as a PNG image to include in presentations or reports. Charts are saved to the database, so they persist between sessions.')

# 9. System Implementation
doc.add_page_break()
doc.add_heading('9. System Implementation', level=1)

doc.add_heading('Technologies Used', level=2)
doc.add_paragraph('The entire system is written in Python 3.11. The web interface uses Streamlit, which is a Python framework specifically designed for building data applications. It handles all the web server details automatically so the developer can focus on the data logic. For data manipulation, the system uses Pandas and NumPy, which are the standard libraries for working with tabular data and numeric operations. Statistical analysis uses SciPy and scikit-learn. Interactive charts are built with Plotly. The database is MySQL 5.7 or later, and the system uses SQLAlchemy as the object-relational mapper so that database operations feel like regular Python code rather than SQL queries.')

doc.add_heading('How It\'s Organized', level=2)
doc.add_paragraph('The main file is app.py, which contains the web interface and the logic for all six sections. It\'s about 618 lines of code. The modules folder contains supporting code organized by function: the profiler module handles statistical analysis, the cleaner module contains the six cleaning operations, the insights module generates findings, the visualizer module builds charts, the db module handles database connections, and the repository module provides functions for saving and loading data. The database folder contains the SQL schema that defines the five tables used to store datasets, column information, cleaning actions, insights, and dashboard configurations. The sample_data folder contains an example employee dataset that has intentional quality issues so users can practice cleaning.')

# 10. Sample Dataset
doc.add_page_break()
doc.add_heading('10. Sample Dataset & Output', level=1)

doc.add_heading('What the Sample Data Contains', level=2)
doc.add_paragraph('The system includes an example dataset called messy_employees.csv with 215 rows and 8 columns. Each row represents an employee with fields for ID, name, department, salary, hire date, employment status, email, and years of experience. This dataset intentionally includes most common data quality problems so that users can practice cleaning without having to create their own messy data.')

doc.add_heading('Problems in the Sample Data', level=2)
doc.add_paragraph('About seven percent of the salary values are missing. The hire date column has four different date formats mixed together: some dates are formatted as YYYY-MM-DD, others as MM/DD/YYYY, some as "Month DD YYYY", and one row contains "not-a-date" which isn\'t a valid date at all. Department names appear in different cases: some as "SALES", some as "Sales", some as "sales". The employment status field has the same capitalization problem. Some name fields have extra spaces at the beginning or end. There are three rows that are exact duplicates of other rows. One salary value is nine million nine hundred ninety-nine thousand nine hundred ninety-nine dollars, which is clearly an error in a dataset where typical salaries range from twenty-eight thousand to one hundred seventy-eight thousand. One employee has negative two years of experience, which is impossible. Overall, about thirty-six cells out of the total, or about four point six percent, contain some kind of data quality issue.')

doc.add_heading('Results After Cleaning', level=2)
doc.add_paragraph('After applying the cleaning operations, the dataset goes from 215 rows with thirty-six quality issues to 209 rows with zero quality issues. Seven salary values are filled using the median. Twelve employment status values are filled using the mode. All departments and statuses are standardized to uppercase. All hire dates are converted to the standard YYYY-MM-DD format. Six rows containing invalid or impossible values are removed. All three duplicate rows are deleted. When you compare the original to the cleaned version, every column is now in a consistent format, all values fall within realistic ranges, and there are no missing values to work around.')

# 11. Limitations
doc.add_page_break()
doc.add_heading('11. Limitations', level=1)

doc.add_heading('What the System Can\'t Do', level=2)
doc.add_paragraph('The system has some practical limitations to be aware of. The file size limit is roughly one hundred megabytes, depending on available RAM. The system isn\'t designed for binary data like images, video, or audio files. Correlation analysis only works for numeric columns; you can\'t correlate text columns against each other. Very large datasets with fifty thousand or more rows will slow down the profiling step. There\'s no multi-user collaboration built in, so each session is isolated to a single user. The system currently only supports MySQL for the database, not PostgreSQL or other alternatives.')

doc.add_paragraph('Some types of data problems can\'t be handled automatically. The system can\'t recognize encrypted or password-protected files. It can\'t detect logical inconsistencies like an employee with a hire date in the future. It doesn\'t use machine learning to predict missing values or to detect anomalies. It can\'t handle unstructured data like text documents or images. It\'s batch-oriented, so it doesn\'t support real-time streaming data. And it can\'t validate data against complex business rules that would require domain expertise.')

# 12. Recommendations
doc.add_page_break()
doc.add_heading('12. Recommendations / Future Improvements', level=1)

doc.add_heading('Near-Term Improvements', level=2)
doc.add_paragraph('In the near term, the most valuable additions would be advanced imputation methods that can predict missing values using machine learning approaches like K-Nearest Neighbors or multiple imputation. A custom validation rules engine would let users define business-specific rules beyond simple range checks. An undo and redo feature with a visual timeline would make it safer to experiment with different cleaning approaches. The ability to export cleaned data to CSV or Excel and to generate PDF reports would integrate better with existing workflows.')

doc.add_heading('Medium-Term Improvements', level=2)
doc.add_paragraph('Looking ahead, multi-file support would let users work with datasets that span multiple CSV files and define relationships between them through joins. Collaborative features including multi-user sessions and shared dashboards would make this useful for teams rather than individuals. More advanced analytics like anomaly detection algorithms, clustering, and time-series analysis would provide deeper insights. And integrating machine learning capabilities would let users build predictive models directly from the cleaned data.')

doc.add_heading('Long-Term Vision', level=2)
doc.add_paragraph('The longer-term vision might include a mobile app for basic exploration on the go, a REST API that would let other systems trigger cleaning workflows automatically, and data lineage tracking that shows where each value came from and what transformations it went through. There\'s also potential to integrate with other data platforms and cloud services, though this would require significant additional development.')

# Summary
doc.add_page_break()
doc.add_heading('Summary', level=1)

doc.add_paragraph('Refine Data Studio is a practical tool designed to solve a real problem that analysts and engineers face every day. Raw data is messy. Cleaning it is necessary but tedious. The system automates the detection and remediation of common data quality problems, provides a clear interface for understanding what\'s wrong and what\'s been fixed, and offers straightforward ways to explore the results. The goal is to compress the time spent on data preparation from hours to minutes, freeing up time for the actual analysis and decision-making.')

doc.add_paragraph()
doc.add_paragraph('The system is built on mature, well-established technologies. The code is organized in a way that makes it easy to understand and extend. The database design allows cleaning workflows and results to be saved and reproduced. And the interface is designed to be intuitive enough that users don\'t need to be programmers to use it effectively.')

doc.add_paragraph()
doc.add_paragraph('Whether you\'re a data analyst working with databases, a business analyst exploring trends, a database administrator validating imports, or a student learning data techniques, the system provides tools that are practical and useful. The sample dataset with intentional problems lets users learn by doing. The six cleaning operations cover the most common scenarios. The automatic insights provide quick wins. And the interactive visualizations make it easy to spot patterns and share results.')

# Save
output_path = 'documents/FEDM_Documentation_Final.docx'
doc.save(output_path)
print(f"SUCCESS: Created {output_path}")
print("Format: Word (.docx)")
print("Style: Natural paragraph format, no bullet points")
print("Pages: 12-15")
